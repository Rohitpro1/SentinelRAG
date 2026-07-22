from __future__ import annotations

import io
import zipfile
import xml.etree.ElementTree as ET
from app.core.exceptions import IngestionError
from app.services.ingestion.parsers.base import BaseDocumentParser, ParsedDocument


class DOCXDocumentParser(BaseDocumentParser):
    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        if not content or len(content) < 10:
            raise IngestionError(f"File {filename} is empty or corrupted.")

        paragraphs: list[str] = []

        try:
            # Try python-docx if installed
            import docx  # type: ignore[import-not-found]
            doc = docx.Document(io.BytesIO(content))
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
            for t in doc.tables:
                for row in t.rows:
                    row_txt = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_txt:
                        paragraphs.append(row_txt)
        except ImportError:
            # Native zipfile + XML extraction fallback for DOCX
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    if "word/document.xml" in zf.namelist():
                        xml_content = zf.read("word/document.xml")
                        tree = ET.fromstring(xml_content)
                        # Extract all text elements (<w:t>)
                        texts = [elem.text for elem in tree.iter() if elem.tag.endswith("t") and elem.text]
                        if texts:
                            paragraphs.append(" ".join(texts))
            except Exception as exc:
                raise IngestionError(f"Corrupted DOCX file {filename}: {exc}") from exc
        except Exception as exc:
            raise IngestionError(f"Corrupted DOCX file {filename}: {exc}") from exc

        full_text = "\n\n".join(paragraphs).strip()
        if not full_text:
            raise IngestionError(f"DOCX file {filename} contains no extractable text.")

        return ParsedDocument(
            text=full_text,
            page_count=1,
            sections=["main"],
            metadata={"filename": filename, "format": "docx"},
        )
